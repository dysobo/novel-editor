import re
from app.core.database import Database


def _get_all_chapters_ordered(db: Database) -> list:
    """递归获取所有章节，按树形顺序展平"""
    result = []
    _flatten(db, 0, result, 0)
    return result


def _flatten(db, parent_id, result, depth):
    chapters = db.get_chapters(parent_id)
    for ch in chapters:
        result.append((dict(ch), depth))
        _flatten(db, ch["id"], result, depth + 1)


def _strip_html(html: str) -> str:
    """简单去除 HTML 标签"""
    text = re.sub(r'<br\s*/?>', '\n', html)
    text = re.sub(r'<p[^>]*>', '', text)
    text = re.sub(r'</p>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&amp;', '&')
    return text.strip()


def export_txt(db: Database, output_path: str):
    """导出为 TXT 文件"""
    chapters = _get_all_chapters_ordered(db)
    with open(output_path, "w", encoding="utf-8") as f:
        for ch, depth in chapters:
            title = ch["title"]
            if depth == 0:
                f.write(f"\n{'=' * 40}\n{title}\n{'=' * 40}\n\n")
            else:
                f.write(f"\n{title}\n{'-' * 20}\n\n")
            content = _strip_html(ch.get("content", "") or "")
            if content:
                f.write(content + "\n\n")


def export_docx(db: Database, output_path: str):
    """导出为 DOCX 文件"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(12)

    chapters = _get_all_chapters_ordered(db)
    for ch, depth in chapters:
        title = ch["title"]
        if depth == 0:
            doc.add_heading(title, level=1)
        else:
            doc.add_heading(title, level=2)
        content = _strip_html(ch.get("content", "") or "")
        if content:
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Cm(0.74)

    doc.save(output_path)
